from docx import Document
#Reading docx file
def reading_docx(filename):
    doc=Document(filename)
    text=[]

    # Iterate over each paragraph in the document
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    # Join all paragraphs into a single string with newlines
    print('\n'.join(text))

    #Adding new Paragraphs

    doc.add_paragraph("This is newly added paragraph")
    doc.save(filename)



reading_docx('sample.docx')
